import argparse
import docx
import os


class Src:
    def __init__(self, filename):
        self.filename = filename
        try:
            f = open(filename, "r", encoding="utf-8")
            self.content = f.read()
            f.close()
        except (PermissionError, UnicodeDecodeError):
            self.content = None


def makeTable(document, src):
    table = document.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = src.filename
    try:
        table.rows[1].cells[0].text = src.content
    except ValueError:
        pass
    table.style = "Table Grid"


class Src2Docx:
    def __init__(self, directory, output):
        self.directory = directory
        self.output = output
        self.srcs = None
        self.searchSrc()

    def run(self):
        document = docx.Document()
        for src in self.srcs:
            makeTable(document, src)
            document.add_paragraph()
        document.save(self.output)

    def searchSrc(self):
        srcs = []
        cwd = os.getcwd()
        for (path, dirs, files) in os.walk(self.directory):
            try:
                os.chdir(path)
            except PermissionError:
                continue
            for filename in files:
                f = Src(str(filename))
                if f.content is not None:
                    print(f.filename)
                    srcs.append(f)
        os.chdir(cwd)
        self.srcs = srcs


if __name__ == "__main__":
    argumentParser = argparse.ArgumentParser()
    argumentParser.add_argument("directory", type=str, help="directory")
    argumentParser.add_argument("output", type=str, help="output")
    args = argumentParser.parse_args()
    src2Docx = Src2Docx(args.directory, args.output)
    src2Docx.run()
